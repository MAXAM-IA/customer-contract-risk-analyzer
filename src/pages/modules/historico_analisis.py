import streamlit as st
import sqlite3
import pandas as pd
import json
import html
import requests
import time
import os
from pathlib import Path
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import io
import re

def obtener_analisis_completados_backend():
    """Obtiene análisis completados desde el backend FastAPI"""
    try:
        # Timeout bajo para evitar bloquear la UI si el backend no responde
        response = requests.get("http://localhost:8000/procesos", timeout=1)
        if response.status_code == 200:
            procesos = response.json()
            # Filtrar solo análisis completados
            completados = [
                p for p in procesos 
                if p.get("estado") == "completado" and p.get("num_resultados", 0) > 0
            ]
            return completados
        else:
            st.error(f"Error al obtener análisis del backend: {response.status_code}")
            return []
    except Exception as e:
        st.error(f"Error de conexión con el backend: {e}")
        return []


def _obtener_nombre_desde_db(analisis_id: str) -> str:
    db_path = Path(__file__).parent.parent.parent / "analisis.db"
    try:
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("SELECT filename FROM analisis WHERE id=?", (analisis_id,))
        row = c.fetchone()
        conn.close()
        if row and row[0]:
            return row[0]
    except Exception:
        return f"Análisis_{analisis_id[:8]}"
    return f"Análisis_{analisis_id[:8]}"

def obtener_analisis_completados(filtro_nombre="", fecha_desde=None, fecha_hasta=None):
    """Obtiene análisis completados con filtros opcionales - PRIMERO INTENTA BACKEND"""
    # Intentar obtener desde el backend primero
    analisis_backend = obtener_analisis_completados_backend()
    if analisis_backend:
        # Convertir formato del backend al formato esperado
        resultados = []
        for proceso in analisis_backend:
            # Extraer información relevante
            id_analisis = proceso.get("id", "")
            nombre_analisis = proceso.get("nombre_analisis")
            if not nombre_analisis:
                nombre_analisis = _obtener_nombre_desde_db(id_analisis)
            created_at = proceso.get("fecha_modificacion", "")
            resultados.append((id_analisis, nombre_analisis, created_at))
        return resultados
    
    # Fallback a SQLite si el backend no está disponible
    conn = sqlite3.connect(Path(__file__).parent.parent.parent / "analisis.db")
    c = conn.cursor()
    
    query = "SELECT id, filename, created_at FROM analisis WHERE estado = '✅ Completed'"
    params = []
    
    if filtro_nombre:
        query += " AND filename LIKE ?"
        params.append(f"%{filtro_nombre}%")
    
    if fecha_desde:
        query += " AND date(created_at) >= ?"
        params.append(fecha_desde.strftime("%Y-%m-%d"))
    
    if fecha_hasta:
        query += " AND date(created_at) <= ?"
        params.append(fecha_hasta.strftime("%Y-%m-%d"))
    
    query += " ORDER BY created_at DESC"
    
    c.execute(query, params)
    filas = c.fetchall()
    conn.close()

    resultados = []
    for analisis_id, filename, created_at in filas:
        nombre_json = None
        try:
            with open(_progreso_path(analisis_id), "r", encoding="utf-8") as f:
                progreso = json.load(f)
            nombre_json = progreso.get("nombre_analisis")
        except Exception:
            nombre_json = None

        nombre = nombre_json or filename or f"Análisis_{analisis_id[:8]}"
        resultados.append((analisis_id, nombre, created_at))

    return resultados

# =============================
# Cachés ligeras para progreso/Word
# =============================

def _progreso_path(analisis_id: str) -> Path:
    return Path(__file__).parent.parent.parent.parent / "fastapi_backend" / "progreso" / f"{analisis_id}.json"

def _progreso_mtime(analisis_id: str) -> float:
    try:
        return os.path.getmtime(_progreso_path(analisis_id))
    except Exception:
        return 0.0

@st.cache_data(show_spinner=False)
def _leer_progreso_cached(analisis_id: str, mtime: float):
    p = _progreso_path(analisis_id)
    if not p.exists():
        return None
    with open(p, "r", encoding="utf-8") as f:
        return json.load(f)

@st.cache_data(show_spinner=False)
def _generar_word_cached(analisis_id: str, filename: str, mtime: float):
    """Genera el Word y lo cachea por ID + timestamp del archivo de progreso."""
    progreso_data = _leer_progreso_cached(analisis_id, mtime)
    if not progreso_data:
        return None
    preguntas = progreso_data.get('preguntas_originales') or progreso_data.get('resultados', [])
    resultados = progreso_data.get('resultados', [])
    return generar_documento_word_corporativo(progreso_data, preguntas, resultados, filename)

def exportar_historico_csv():
    """Exporta el histórico completo a CSV"""
    try:
        conn = sqlite3.connect(Path(__file__).parent.parent.parent / "analisis.db")
        df = pd.read_sql_query(
            "SELECT id, filename, estado, created_at FROM analisis ORDER BY created_at DESC", 
            conn
        )
        conn.close()
        
        # Convertir a CSV
        csv = df.to_csv(index=False)
        return csv
    except Exception as e:
        st.error(f"Error al exportar: {str(e)}")
        return None

@st.dialog("Detalle de análisis", width="large")
def mostrar_detalle_dialog(analisis_id, filename):
    """Dialog para mostrar el detalle de un análisis completado"""
    progreso_path = Path(__file__).parent.parent.parent.parent / "fastapi_backend" / "progreso" / f"{analisis_id}.json"
    
    # Cabecera del modal más compacta con botón de refrescar
    col_header1, col_header2 = st.columns([3, 1], vertical_alignment="center", border=False)
    
    with col_header1:
        st.markdown(f"""
            <div style='
                background: linear-gradient(90deg, #f9fafb 0%, #f3f4f6 100%);
                border-radius: 6px;
                padding: 0.75rem 1rem;
                margin-bottom: 1rem;
                border-left: 3px solid #dc2626;
            '>
                <h3 style='
                    font-size: 1rem;
                    font-weight: 600;
                    color: #1f2937;
                    margin: 0 0 0.25rem 0;
                '>{filename}</h3>
                <div style='
                    color: #6b7280;
                    font-size: 0.8rem;
                    display: flex;
                    gap: 1rem;
                '>
                    <span>🆔 {analisis_id[:8]}...</span>
                    <span>✅ Análisis Completado</span>
                </div>
            </div>
        """, unsafe_allow_html=True)
    
    with col_header2:
        if st.button("🔄 Refrescar", help="Actualizar datos del análisis", use_container_width=True):
            # Forzar recarga limpiando cualquier cache
            if f'progreso_data_{analisis_id}' in st.session_state:
                del st.session_state[f'progreso_data_{analisis_id}']
            st.rerun()
    
    if progreso_path.exists():
        try:
            # Auto-reload si está activado
            if st.session_state.get(f'auto_reload_{analisis_id}', False):
                # Marcar tiempo de auto-reload
                if f'last_reload_{analisis_id}' not in st.session_state:
                    st.session_state[f'last_reload_{analisis_id}'] = time.time()
                
                # Auto-reload cada 3 segundos
                if time.time() - st.session_state[f'last_reload_{analisis_id}'] > 3:
                    st.session_state[f'last_reload_{analisis_id}'] = time.time()
                    st.rerun()
                
                # Mostrar indicador de auto-reload
                st.markdown("""
                    <div style='
                        background: #dbeafe;
                        border: 1px solid #93c5fd;
                        border-radius: 6px;
                        padding: 0.5rem;
                        margin-bottom: 1rem;
                        text-align: center;
                        color: #1d4ed8;
                        font-size: 0.8rem;
                    '>
                        🔄 Auto-reload activado - Actualizando cada 3 segundos
                    </div>
                """, unsafe_allow_html=True)
            
            # Leer datos SIEMPRE fresco del archivo (no usar cache)
            with open(progreso_path, "r", encoding="utf-8") as f:
                progreso_data = json.load(f)
            
            # Detectar si hay reanalisis recientes y cambios
            fecha_modificacion = progreso_data.get('fecha_modificacion', '')
            tipo_reanalisis = progreso_data.get('tipo_reanalisis', '')
            
            # Verificar si hay preguntas reanalazadas en los resultados
            preguntas_reanalazadas = []
            for idx, resultado in enumerate(progreso_data.get('resultados', [])):
                if resultado.get('reanalizado_en'):
                    preguntas_reanalazadas.append(idx + 1)
            
            # Mostrar alertas mejoradas si hay cambios recientes
            if preguntas_reanalazadas:
                st.markdown(f"""
                    <div style='
                        background: #ecfdf5;
                        border: 1px solid #10b981;
                        border-left: 4px solid #10b981;
                        border-radius: 5px;
                        padding: 0.25rem 0.5rem;
                        margin-bottom: 0.4rem;
                        width: 100%;
                        max-width: 100vw;
                    '>
                        <div style='
                            display: flex;
                            align-items: center;
                            gap: 0.3rem;
                            color: #064e3b;
                            font-weight: 600;
                            font-size: 0.85rem;
                        '>
                            🔄 <strong>Cambios Detectados</strong>
                        </div>
                        <div style='color: #065f46; font-size: 0.75rem;'>
                            {"Pregunta" if len(preguntas_reanalazadas) == 1 else "Preguntas"} 
                            #{", #".join(map(str, preguntas_reanalazadas))} 
                            {"ha sido reanalizada" if len(preguntas_reanalazadas) == 1 else "han sido reanalizadas"}
                        </div>
                        <div style='color: #047857; font-size: 0.7rem; margin-top: 0.05rem;'>
                            📅 Última modificación: {fecha_modificacion}
                        </div>
                    </div>
                """, unsafe_allow_html=True)
            elif tipo_reanalisis and 'global' in tipo_reanalisis.lower():
                st.info(f"🔄 **Reanalisis global detectado**: Todo el documento fue reprocesado. Última modificación: {fecha_modificacion}")
            
            preguntas = progreso_data.get('preguntas_originales') or progreso_data.get('resultados', [])
            resultados = progreso_data.get('resultados', [])
            
            # Resumen ejecutivo compacto
            total_preguntas = len(preguntas)
            completadas = sum(1 for r in resultados if r.get('Estado') == '✅ Completed')
            
            # Estadísticas de riesgo - versión mejorada y consistente
            riesgos = {'Alto': 0, 'Medio': 0, 'Bajo': 0, 'Sin evaluar': 0}
            
            # Contar riesgos de manera consistente
            for i, resultado in enumerate(resultados):
                # Buscar riesgo en diferentes campos posibles (Riesgo, RISK, risk)
                riesgo = (resultado.get('Riesgo') or 
                         resultado.get('RISK') or 
                         resultado.get('risk') or 
                         'Sin evaluar').strip()
                
                # Normalizar valores de riesgo en inglés y español
                if riesgo.upper() in ['ALTO', 'HIGH']:
                    riesgos['Alto'] += 1
                elif riesgo.upper() in ['MEDIO', 'MEDIUM']:
                    riesgos['Medio'] += 1
                elif riesgo.upper() in ['BAJO', 'LOW']:
                    riesgos['Bajo'] += 1
                else:
                    riesgos['Sin evaluar'] += 1
            
            # Debug info opcional (más discreto)
            #with st.expander("🔧 Debug Riesgos", expanded=False):
            #    st.write(f"� **Conteos de riesgo:**")
            #    for nivel, cantidad in riesgos.items():
            #        porcentaje = (cantidad/len(resultados)*100) if len(resultados) > 0 else 0
            #        st.write(f"  - {nivel}: {cantidad} ({porcentaje:.1f}%)")
            #    
            #    st.write(f"📋 **Detalle por pregunta:**")
            #    for i, r in enumerate(resultados):
            #        riesgo_raw = r.get('Riesgo', 'Sin evaluar')
            #        st.write(f"  P{i+1}: '{riesgo_raw}'")
            
            col_stats1, col_stats2, col_stats3, col_stats4, col_stats5, col_stats6 = st.columns(6)
            
            with col_stats1:
                st.metric("📋 Preguntas", total_preguntas)
            
            with col_stats2:
                st.metric("✅ Completadas", completadas)
            
            with col_stats3:
                st.metric("🔴 Alto", riesgos['Alto'], delta=f"{riesgos['Alto']/total_preguntas*100:.0f}%" if total_preguntas > 0 else "0%")
            
            with col_stats4:
                st.metric("🟡 Medio", riesgos['Medio'], delta=f"{riesgos['Medio']/total_preguntas*100:.0f}%" if total_preguntas > 0 else "0%")
            
            with col_stats5:
                st.metric("🟢 Bajo", riesgos['Bajo'], delta=f"{riesgos['Bajo']/total_preguntas*100:.0f}%" if total_preguntas > 0 else "0%")
            
            with col_stats6:
                st.metric("⚪ Sin evaluar", riesgos['Sin evaluar'], delta=f"{riesgos['Sin evaluar']/total_preguntas*100:.0f}%" if total_preguntas > 0 else "0%")
            
            # Botón de debug para riesgos
            #if st.button("🔧 Debug Riesgos", help="Mostrar información detallada de evaluación de riesgos"):
            #    st.session_state['show_debug_riesgos'] = not st.session_state.get('show_debug_riesgos', False)
            #    if st.session_state['show_debug_riesgos']:
            #        st.success("✅ Debug de riesgos activado - se mostrará información detallada en cada pregunta")
            #    else:
            #        st.info("ℹ️ Debug de riesgos desactivado")
            #    st.rerun()
            
            # Sección de preguntas y respuestas más compacta
            st.markdown("### 📋 Preguntas y Respuestas")
            
            # Filtros en una sola línea
            col_filter1, col_filter2, col_filter3 = st.columns([2, 2, 1])
            
            with col_filter1:
                secciones = list(set([p.get('Sección', 'Sin sección') for p in preguntas]))
                seccion_filtro = st.selectbox("🔍 Filtrar por sección:", ["Todas"] + secciones, key="seccion_filter")
            
            with col_filter2:
                tipo_filtro = st.selectbox("📊 Mostrar:", [
                    "Todas las respuestas",
                    "Solo riesgos altos", 
                    "Solo riesgos medios",
                    "Solo riesgos bajos",
                    "Con riesgos detectados"
                ], key="tipo_filter")
            
            with col_filter3:
                st.markdown("<div style='margin-top: 1.5rem;'></div>", unsafe_allow_html=True)
                expandir_todo = st.checkbox("📖 Expandir todo", key="expandir_filter")
            
            # Mostrar preguntas filtradas
            for idx, preg in enumerate(preguntas):
                if idx >= len(resultados):
                    continue
                
                respuesta_data = resultados[idx]
                
                # IMPORTANTE: Si la pregunta fue reanalizada, usar la pregunta del resultado
                # en lugar de la pregunta original
                if respuesta_data.get('reanalizado_en') and respuesta_data.get('Pregunta'):
                    # Usar la pregunta reanalizada del resultado
                    pregunta = respuesta_data.get('Pregunta', '')
                    seccion = respuesta_data.get('Sección', preg.get('Sección', 'Sin sección'))
                else:
                    # Usar la pregunta original
                    pregunta = preg.get('Pregunta', '')
                    seccion = preg.get('Sección', 'Sin sección')
                
                respuesta = respuesta_data.get('Respuesta', '')
                
                # Aplicar filtros
                if seccion_filtro != "Todas" and seccion != seccion_filtro:
                    continue
                
                # Aplicar filtro de tipo de riesgo - usar la misma lógica de normalización
                riesgo_nivel_raw = (respuesta_data.get('Riesgo') or 
                                  respuesta_data.get('RISK') or 
                                  respuesta_data.get('risk') or 
                                  'Sin evaluar').strip()
                
                # Normalizar valores de riesgo para filtros
                if riesgo_nivel_raw.upper() in ['ALTO', 'HIGH']:
                    riesgo_nivel = 'Alto'
                elif riesgo_nivel_raw.upper() in ['MEDIO', 'MEDIUM']:
                    riesgo_nivel = 'Medio'
                elif riesgo_nivel_raw.upper() in ['BAJO', 'LOW']:
                    riesgo_nivel = 'Bajo'
                else:
                    riesgo_nivel = 'Sin evaluar'
                
                if tipo_filtro == "Solo riesgos altos" and riesgo_nivel != "Alto":
                    continue
                elif tipo_filtro == "Solo riesgos medios" and riesgo_nivel != "Medio":
                    continue
                elif tipo_filtro == "Solo riesgos bajos" and riesgo_nivel != "Bajo":
                    continue
                elif tipo_filtro == "Con riesgos detectados" and not any(word in respuesta.lower() for word in ['riesgo', 'peligro', 'problema', 'advertencia', 'cuidado']):
                    continue
                
                # Determinar color basado en contenido de la respuesta
                if any(word in respuesta.lower() for word in ['riesgo', 'peligro', 'problema']):
                    border_color = "#dc2626"
                    bg_color = "#fef2f2"
                elif any(word in respuesta.lower() for word in ['favorable', 'beneficio', 'positivo']):
                    border_color = "#16a34a"
                    bg_color = "#f0fdf4"
                else:
                    border_color = "#6b7280"
                    bg_color = "#f9fafb"
                
                with st.expander(f"❓ {seccion} - Pregunta {idx+1}", expanded=expandir_todo):
                    # Verificar si esta pregunta fue reanalizada
                    respuesta_reanalizada = respuesta_data.get('reanalizado_en')
                    tipo_reanalisis_pregunta = respuesta_data.get('tipo_reanalisis')
                    
                    # Mostrar indicador de reanalisis si aplica
                    if respuesta_reanalizada:
                        st.markdown(f"""
                            <div style='
                                background: #fefce8;
                                border: 1px solid #facc15;
                                border-radius: 6px;
                                padding: 0.5rem;
                                margin-bottom: 0.75rem;
                            '>
                                <div style='
                                    display: flex;
                                    align-items: center;
                                    gap: 0.5rem;
                                    color: #854d0e;
                                    font-size: 0.8rem;
                                    font-weight: 600;
                                '>
                                    🔄 <strong>Pregunta Reanalizada</strong>
                                    <span style='
                                        background: #fef3c7;
                                        padding: 0.2rem 0.4rem;
                                        border-radius: 4px;
                                        font-size: 0.7rem;
                                        margin-left: auto;
                                    '>{respuesta_reanalizada}</span>
                                </div>
                            </div>
                        """, unsafe_allow_html=True)
                    
                    # Información básica compacta
                    col1, col2 = st.columns([3, 1])
                    
                    with col1:
                        # Mostrar pregunta con indicador visual si fue modificada
                        if respuesta_reanalizada and respuesta_data.get('Pregunta') != preg.get('Pregunta'):
                            st.markdown(f"**Pregunta (Modificada):** {pregunta}")
                            
                            # Mostrar la pregunta original con un desplegable HTML
                            st.markdown(f"""
                                <details style='margin-top: 0.5rem;'>
                                    <summary style='
                                        color: #6b7280; 
                                        font-size: 0.8rem; 
                                        cursor: pointer;
                                        font-weight: 500;
                                        padding: 0.25rem 0;
                                    '>
                                        👁️ Ver pregunta original
                                    </summary>
                                    <div style='
                                        background: #f9fafb;
                                        border-left: 2px solid #d1d5db;
                                        padding: 0.5rem;
                                        margin: 0.25rem 0;
                                        font-style: italic;
                                        color: #6b7280;
                                        font-size: 0.9rem;
                                        border-radius: 4px;
                                    '>
                                        {preg.get('Pregunta', 'N/A')}
                                    </div>
                                </details>
                            """, unsafe_allow_html=True)
                        else:
                            st.markdown(f"**Pregunta:** {pregunta}")
                    
                    with col2:
                        # Detectar y mostrar nivel de riesgo usando la misma lógica que en las estadísticas
                        riesgo_nivel = (respuesta_data.get('Riesgo') or 
                                      respuesta_data.get('RISK') or 
                                      respuesta_data.get('risk') or 
                                      'Sin evaluar').strip()
                        
                        # Normalizar valores de riesgo en inglés y español
                        if riesgo_nivel.upper() in ['ALTO', 'HIGH']:
                            riesgo_nivel = 'Alto'
                        elif riesgo_nivel.upper() in ['MEDIO', 'MEDIUM']:
                            riesgo_nivel = 'Medio'
                        elif riesgo_nivel.upper() in ['BAJO', 'LOW']:
                            riesgo_nivel = 'Bajo'
                        else:
                            riesgo_nivel = 'Sin evaluar'
                        
                        # Debug temporal: mostrar información del riesgo
                        if st.session_state.get('show_debug_riesgos', False):
                            st.write(f"🔧 Debug P{idx+1}: '{riesgo_nivel}' (Keys: {list(respuesta_data.keys())})")
                            st.write(f"🔧 Raw data: Riesgo={respuesta_data.get('Riesgo')}, RISK={respuesta_data.get('RISK')}, risk={respuesta_data.get('risk')}")
                        
                        if riesgo_nivel == 'Alto':
                            color_riesgo = "#dc2626"
                            emoji_riesgo = "🔴"
                        elif riesgo_nivel == 'Medio':
                            color_riesgo = "#f59e0b"
                            emoji_riesgo = "🟡"
                        elif riesgo_nivel == 'Bajo':
                            color_riesgo = "#16a34a"
                            emoji_riesgo = "🟢"
                        else:
                            color_riesgo = "#6b7280"
                            emoji_riesgo = "⚪"
                        
                        st.markdown(f"""
                            <div style='
                                background: {color_riesgo}15;
                                color: {color_riesgo};
                                padding: 0.25rem 0.5rem;
                                border-radius: 12px;
                                text-align: center;
                                font-size: 0.8rem;
                                font-weight: 600;
                                border: 1px solid {color_riesgo}40;
                            '>
                                {emoji_riesgo} {riesgo_nivel}
                            </div>
                        """, unsafe_allow_html=True)
                    
                    # Pestañas para organizar mejor el contenido
                    #tab1, tab2, tab3 = st.tabs(["📖 Respuesta", "✏️ Editar", "📊 Info"])
                    tab1, tab2 = st.tabs(["📖 Respuesta", "✏️ Editar"])

                    with tab1:
                        # Respuesta de forma compacta
                        st.write(respuesta)
                        
                        # Debug info como desplegable HTML (sin expander anidado)
                        pregunta_original = preg.get('Pregunta', 'N/A')
                        pregunta_resultado = respuesta_data.get('Pregunta', 'N/A')
                        
                        debug_info = {
                            "indice": idx,
                            "pregunta_mostrada": pregunta,
                            "pregunta_original": pregunta_original,
                            "pregunta_en_resultado": pregunta_resultado,
                            "seccion_mostrada": seccion,
                            "seccion_original": preg.get('Sección', 'N/A'),
                            "seccion_en_resultado": respuesta_data.get('Sección', 'N/A'),
                            "respuesta_length": len(respuesta),
                            "riesgo": respuesta_data.get('Riesgo', 'N/A'),
                            "reanalizado_en": respuesta_data.get('reanalizado_en', 'No'),
                            "tipo_reanalisis": respuesta_data.get('tipo_reanalisis', 'N/A'),
                            "fue_reanalizada": bool(respuesta_data.get('reanalizado_en')),
                            "todas_keys_resultado": list(respuesta_data.keys()),
                            "todas_keys_original": list(preg.keys())
                        }
                        
                        debug_json = json.dumps(debug_info, indent=2, ensure_ascii=False)
                        
                        st.markdown(f"""
                            <details style='margin-top: 1rem;'>
                                <summary style='
                                    color: #6b7280; 
                                    font-size: 0.8rem; 
                                    cursor: pointer;
                                    font-weight: 600;
                                    padding: 0.25rem 0;
                                '>
                                    🔧 Debug Info
                                </summary>
                                <div style='
                                    background: #f9fafb;
                                    border: 1px solid #e5e7eb;
                                    border-radius: 6px;
                                    padding: 0.75rem;
                                    margin-top: 0.5rem;
                                    font-family: monospace;
                                    font-size: 0.75rem;
                                    color: #374151;
                                    overflow-x: auto;
                                '>
                                    <pre>{debug_json}</pre>
                                </div>
                            </details>
                        """, unsafe_allow_html=True)
                    
                    with tab2:
                        # Formulario de edición
                        st.markdown("**Editar pregunta para re-análisis:**")
                        
                        nueva_pregunta = st.text_area(
                            "Pregunta modificada:",
                            value=pregunta,
                            height=80,
                            key=f"edit_pregunta_{analisis_id}_{idx}",
                            help="Edita la pregunta y presiona 'Re-analizar' para generar una nueva respuesta"
                        )
                        
                        nueva_seccion = st.text_input(
                            "Sección:",
                            value=seccion,
                            key=f"edit_seccion_{analisis_id}_{idx}",
                            help="Opcional: editar la sección de la pregunta"
                        )
                        
                        col_btn1, col_btn2 = st.columns([1, 1])
                        
                        with col_btn1:
                            # Mostrar estado del botón para debugging
                            #st.write(f"🔧 Debug: ID={analisis_id[:8]}, idx={idx}")
                            
                            # Mostrar timestamp del archivo para verificar actualizaciones
                            if progreso_path.exists():
                                file_mtime = os.path.getmtime(progreso_path)
                                file_time = datetime.fromtimestamp(file_mtime).strftime("%H:%M:%S")
                                #st.write(f"📁 Archivo actualizado: {file_time}")
                            
                            if st.button(
                                "🔄 Re-analizar Pregunta", 
                                key=f"reanalizar_{analisis_id}_{idx}",
                                use_container_width=True,
                                type="primary",
                                help="Hace click para re-analizar esta pregunta individual"
                            ):
                                # Re-análisis asíncrono - enviar al backend y mostrar en procesos en curso
                                from db.analisis_db import guardar_analisis
                                
                                #st.write("🚀 ¡Botón presionado! Iniciando re-análisis...")
                                #st.write(f"📝 Pregunta original: {pregunta[:50]}...")
                                #st.write(f"📝 Pregunta editada: {nueva_pregunta[:50]}...")
                                
                                try:
                                    API_URL = "http://localhost:8000"  # Ajustar según tu configuración
                                    payload = {
                                        "pregunta": nueva_pregunta,  # Usar la pregunta editada
                                        "seccion": nueva_seccion     # Usar la sección editada
                                    }
                                    
                                    #st.write(f"📤 Enviando: {payload}")
                                    
                                    # PRIMERO: Registrar en BD como proceso pendiente
                                    filename_proceso = f"Reanalisis_P{idx+1}_{filename[:20]}..."
                                    guardar_analisis(
                                        id=analisis_id,  # Usar el mismo ID para sobreescribir
                                        filename=filename_proceso,
                                        estado="🔄 Reprocesando"
                                    )
                                    #st.write(f"✅ Proceso registrado en BD como: {filename_proceso}")
                                    
                                    # SEGUNDO: Enviar al backend
                                    with st.spinner("Enviando re-análisis al backend..."):
                                        response = requests.post(
                                            f"{API_URL}/reanalisar_pregunta/{analisis_id}/{idx}",
                                            json=payload,
                                            timeout=10  # Timeout corto para envío
                                        )
                                    
                                    #st.write(f"📥 Respuesta del servidor: {response.status_code}")
                                    
                                    if response.status_code == 200:
                                        result = response.json()
                                        proceso_id = result.get("id")
                                        mensaje = result.get("mensaje", "Re-análisis iniciado")
                                        st.toast(f"Re-analizando pregunta", icon="🔄")
                                        #st.success(f"✅ {mensaje}")
                                        #st.success(f"📝 Usando pregunta editada: '{nueva_pregunta[:50]}...'")
                                        
                                        # Verificar si es el mismo ID (sobreescribiendo) o un nuevo ID
                                        if proceso_id == analisis_id:
                                            pass
                                            #st.info(f"🔄 Re-analizando pregunta #{idx+1} en el análisis actual: {analisis_id[:8]}...")
                                            #st.info("📝 El resultado se actualizará en este mismo análisis.")
                                        else:
                                            #st.info(f"🆔 Nuevo proceso creado: {proceso_id[:8]}...")
                                            # Si es un nuevo ID, también registrarlo
                                            if proceso_id != analisis_id:
                                                guardar_analisis(
                                                    id=proceso_id,
                                                    filename=f"Reanalisis_P{idx+1}_{filename[:20]}...",
                                                    estado="🔄 Reprocesando"
                                                )
                                        
                                        #st.info("🔄 El proceso aparecerá en 'Procesos en Curso'. Los resultados se actualizarán automáticamente.")
                                        
                                        # Mostrar enlace directo a procesos en curso
                                        '''
                                        st.markdown("""
                                            <div style='
                                                background: #ecfdf5;
                                                border: 1px solid #bbf7d0;
                                                border-radius: 6px;
                                                padding: 0.75rem;
                                                margin: 0.5rem 0;
                                            '>
                                                <p style='margin: 0; color: #065f46; font-size: 0.9rem;'>
                                                    ✨ <strong>¡Proceso iniciado!</strong><br>
                                                    📍 Ve a <strong>"Procesos en Curso"</strong> para ver el progreso en tiempo real
                                                </p>
                                            </div>
                                        """, unsafe_allow_html=True)
                                        '''
            
                                            
                                    else:
                                        st.error(f"❌ Error: {response.status_code} - {response.text}")
                                        # Si falla, revertir el estado en BD
                                        from db.analisis_db import actualizar_estado_analisis
                                        actualizar_estado_analisis(analisis_id, "✅ Completed")
                                        
                                except requests.exceptions.ConnectionError:
                                    st.error("❌ No se puede conectar al backend. ¿Está ejecutándose?")
                                    # Revertir estado en BD
                                    from db.analisis_db import actualizar_estado_analisis
                                    actualizar_estado_analisis(analisis_id, "✅ Completed")
                                except requests.exceptions.Timeout:
                                    st.error("❌ Timeout al enviar la solicitud. El backend puede estar ocupado.")
                                    # Revertir estado en BD
                                    from db.analisis_db import actualizar_estado_analisis
                                    actualizar_estado_analisis(analisis_id, "✅ Completed")
                                except Exception as e:
                                    st.error(f"❌ Error inesperado: {str(e)}")
                                    st.write(f"🔍 Detalles del error: {type(e).__name__}: {str(e)}")
                                    # Revertir estado en BD
                                    from db.analisis_db import actualizar_estado_analisis
                                    actualizar_estado_analisis(analisis_id, "✅ Completed")
                        
                        with col_btn2:
                            if st.button(
                                "🔄 Re-analizar Todo", 
                                key=f"reanalizar_todo_{analisis_id}_{idx}",
                                use_container_width=True,
                                help="Re-analizar todo el documento con preguntas modificadas"
                            ):
                                # Re-análisis global asíncrono
                                from db.analisis_db import guardar_analisis
                                
                                try:
                                    # PRIMERO: Registrar en BD como proceso pendiente
                                    filename_proceso = f"Reanalisis_Global_{filename[:15]}..."
                                    guardar_analisis(
                                        id=analisis_id,  # Usar el mismo ID para sobreescribir
                                        filename=filename_proceso,
                                        estado="🔄 Reprocesando Global"
                                    )
                                    st.write(f"✅ Proceso global registrado en BD como: {filename_proceso}")
                                    
                                    # Recopilar todas las preguntas (originales + editadas)
                                    todas_preguntas = []
                                    preguntas_editadas = 0
                                    for i, (p_original, r_data) in enumerate(zip(preguntas, resultados)):
                                        # Usar preguntas editadas si existen en session_state
                                        pregunta_key = f"edit_pregunta_{analisis_id}_{i}"
                                        seccion_key = f"edit_seccion_{analisis_id}_{i}"
                                        
                                        pregunta_final = st.session_state.get(pregunta_key, p_original.get('Pregunta', ''))
                                        seccion_final = st.session_state.get(seccion_key, p_original.get('Sección', ''))
                                        
                                        # Verificar si fue editada
                                        if pregunta_final != p_original.get('Pregunta', '') or seccion_final != p_original.get('Sección', ''):
                                            preguntas_editadas += 1
                                        
                                        todas_preguntas.append({
                                            "pregunta": pregunta_final,
                                            "seccion": seccion_final
                                        })
                                    
                                    st.write(f"📝 Total de preguntas: {len(todas_preguntas)}")
                                    st.write(f"✏️ Preguntas editadas: {preguntas_editadas}")
                                    
                                    API_URL = "http://localhost:8000"
                                    payload = {"preguntas": todas_preguntas}
                                    
                                    with st.spinner("Enviando re-análisis global al backend..."):
                                        response = requests.post(
                                            f"{API_URL}/reanalisar_global/{analisis_id}",
                                            json=payload,
                                            timeout=15  # Timeout más largo para análisis global
                                        )
                                    
                                    if response.status_code == 200:
                                        result = response.json()
                                        proceso_id = result.get("id")
                                        mensaje = result.get("mensaje", "Re-análisis global iniciado")
                                        
                                        st.success(f"✅ {mensaje}")
                                        
                                        # Verificar si es el mismo ID (sobreescribiendo) o un nuevo ID
                                        if proceso_id == analisis_id:
                                            st.info(f"🔄 Re-analizando TODAS las preguntas en el análisis actual: {analisis_id[:8]}...")
                                            st.info("📝 Los resultados se actualizarán en este mismo análisis.")
                                        else:
                                            st.info(f"🆔 Nuevo proceso creado: {proceso_id[:8]}...")
                                        
                                        st.info("🔄 El proceso aparecerá en 'Procesos en Curso'. Los resultados se actualizarán automáticamente.")
                                        
                                        # Opcional: redirigir a procesos en curso
                                        if st.button("👀 Ver en Procesos en Curso (Global)", type="secondary"):
                                            st.switch_page("pages/procesos_en_curso.py")
                                            
                                    else:
                                        st.error(f"❌ Error: {response.status_code} - {response.text}")
                                        
                                except requests.exceptions.ConnectionError:
                                    st.error("❌ No se puede conectar al backend. ¿Está ejecutándose?")
                                except requests.exceptions.Timeout:
                                    st.error("❌ Timeout al enviar la solicitud. El backend puede estar ocupado.")
                                except Exception as e:
                                    st.error(f"❌ Error inesperado: {str(e)}")
                                    
                                st.rerun()  # Refrescar para ver los cambios
                    
                    #with tab3:
                    #    # Información adicional compacta
                    #    info_cols = st.columns(3)
                    #    
                    #    with info_cols[0]:
                    #        st.metric("Pregunta #", f"{idx+1}")
                    #    
                    #    with info_cols[1]:
                    #        st.metric("Sección", seccion)
                    #        
                    #    with info_cols[2]:
                    #        estado = respuesta_data.get('Estado', 'Desconocido')
                    #        st.metric("Estado", estado)
        
        except Exception as e:
            st.error(f"Error al cargar el análisis: {str(e)}")
    else:
        st.warning("⚠️ No se encontró el archivo de progreso para este análisis.")

def mostrar_historico():
    """Función principal para mostrar el histórico de análisis"""
    st.markdown("""
        <h2 style='
            color: #dc2626;
            text-align: center;
            margin-bottom: 0.5rem;
            font-size: 1.2rem;
            font-weight: 600;
        '>📚 Histórico de Análisis</h2>
    """, unsafe_allow_html=True)
    
    # Información introductoria muy compacta
    st.markdown("""
        <div style='
            background: #f0f9ff;
            border: 1px solid #bae6fd;
            border-left: 3px solid #0ea5e9;
            padding: 0.5rem 0.75rem;
            border-radius: 6px;
            margin-bottom: 0.75rem;
        '>
            <div style='
                display: flex;
                align-items: center;
                gap: 0.5rem;
            '>
                <span style='
                    color: #0ea5e9;
                    font-size: 0.9rem;
                '>📚</span>
                <span style='
                    color: #0c4a6e;
                    font-size: 0.85rem;
                    font-weight: 600;
                '>Repositorio de Análisis Completados</span>
                <span style='
                    color: #0369a1;
                    font-size: 0.75rem;
                    margin-left: auto;
                '>Buscar y exportar análisis</span>
            </div>
        </div>
    """, unsafe_allow_html=True)
    
    # Filtros compactos y discretos
    with st.expander("🔍 Filtros", expanded=False):
        # Primera fila con filtros principales - todos sin título para mantener alineación
        col_filtro1, col_filtro2, col_filtro3, col_filtro4 = st.columns([2.5, 1.2, 1.2, 0.8])
        
        with col_filtro1:
            filtro_nombre = st.text_input(
                "Buscar:",
                placeholder="🔍 Buscar por nombre de archivo...",
                help="Buscar por nombre de archivo",
                label_visibility="collapsed",
                key="filtro_nombre_input"
            )
        
        with col_filtro2:
            if 'fecha_desde_historico' not in st.session_state:
                st.session_state.fecha_desde_historico = datetime.now() - timedelta(days=30)
            
            fecha_desde = st.date_input(
                ".",
                value=st.session_state.fecha_desde_historico,
                key="fecha_desde_input",
                label_visibility="collapsed",
                help="Fecha desde"
            )
            st.session_state.fecha_desde_historico = fecha_desde
            
        with col_filtro3:
            if 'fecha_hasta_historico' not in st.session_state:
                st.session_state.fecha_hasta_historico = datetime.now()
            
            fecha_hasta = st.date_input(
                ".",
                value=st.session_state.fecha_hasta_historico,
                key="fecha_hasta_input",
                label_visibility="collapsed",
                help="Fecha hasta"
            )
            st.session_state.fecha_hasta_historico = fecha_hasta
        
        with col_filtro4:
            if st.button("Clear 🧹", help="Limpiar filtros", use_container_width=True):
                if 'fecha_desde_historico' in st.session_state:
                    del st.session_state.fecha_desde_historico
                if 'fecha_hasta_historico' in st.session_state:
                    del st.session_state.fecha_hasta_historico
                if 'filtro_nombre_input' in st.session_state:
                    del st.session_state.filtro_nombre_input
                st.rerun()
    
        # Etiquetas discretas debajo de los campos para claridad
        st.markdown("""
            <div style='
                display: flex;
                justify-content: space-between;
                margin-top: 0.2rem;
                font-size: 0.65rem;
                color: #64748b;
                padding: 0 0.2rem;
            '>
                <span style='flex: 2.3; text-align: left;'>Nombre del archivo</span>
                <span style='flex: 1.2; text-align: center;'>Desde</span>
                <span style='flex: 1.2; text-align: center;'>Hasta</span>
                <span style='flex: 0.8; text-align: center;'>Limpiar</span>
            </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<div style='margin: 1.5rem 0;'></div>", unsafe_allow_html=True)
    
    # Panel de estadísticas después de los filtros
    col_stats = st.columns([1])[0]
    
    with col_stats:
        # Panel de estadísticas después de los filtros
        try:
            conn = sqlite3.connect(Path(__file__).parent.parent.parent / "analisis.db")
            c = conn.cursor()
            
            # Estadísticas básicas
            c.execute("SELECT COUNT(*) FROM analisis WHERE estado = '✅ Completed'")
            total_completados = c.fetchone()[0]
            
            c.execute("SELECT COUNT(*) FROM analisis WHERE date(created_at) = date('now')")
            hoy = c.fetchone()[0]
            
            c.execute("SELECT COUNT(*) FROM analisis WHERE date(created_at) >= date('now', '-7 days')")
            esta_semana = c.fetchone()[0]
            
            conn.close()
            
            st.markdown(f"""
                <div style='
                    background: #ffffff;
                    border: 1px solid #e5e7eb;
                    border-radius: 6px;
                    padding: 0.5rem 0.75rem;
                    box-shadow: 0 1px 4px rgba(0,0,0,0.05);
                    margin-bottom: 0.75rem;
                '>
                    <div style='
                        display: flex;
                        justify-content: space-around;
                        align-items: center;
                        text-align: center;
                    '>
                        <div>
                            <div style='font-size: 1.1rem; font-weight: 700; color: #dc2626;'>{total_completados}</div>
                            <div style='font-size: 0.65rem; color: #6b7280;'>Total</div>
                        </div>
                        <div style='border-left: 1px solid #e5e7eb; height: 30px;'></div>
                        <div>
                            <div style='font-size: 1.1rem; font-weight: 700; color: #16a34a;'>{hoy}</div>
                            <div style='font-size: 0.65rem; color: #6b7280;'>Hoy</div>
                        </div>
                        <div style='border-left: 1px solid #e5e7eb; height: 30px;'></div>
                        <div>
                            <div style='font-size: 1.1rem; font-weight: 700; color: #2563eb;'>{esta_semana}</div>
                            <div style='font-size: 0.65rem; color: #6b7280;'>Esta Semana</div>
                        </div>
                    </div>
                </div>
            """, unsafe_allow_html=True)
            
        except Exception:
            st.error("Error al cargar estadísticas")
    
    st.markdown("<div style='margin: 1rem 0;'></div>", unsafe_allow_html=True)
    
    # Configurar paginación
    items_por_pagina = 10
    if 'pagina_actual_historico' not in st.session_state:
        st.session_state.pagina_actual_historico = 1

    # Intentar backend primero (mantenemos comportamiento: sin filtros en backend)
    resultados_backend = obtener_analisis_completados_backend()

    resultados_pagina = []
    total_items = 0

    if resultados_backend:
        resultados_full = []
        for proceso in resultados_backend:
            id_analisis = proceso.get("id", "")
            filename = proceso.get("nombre_analisis") or proceso.get("filename")
            if not filename:
                filename = _obtener_nombre_desde_db(id_analisis)
            created_at = proceso.get("fecha_modificacion", "")
            resultados_full.append((id_analisis, filename, created_at))

        total_items = len(resultados_full)
        total_paginas = (total_items - 1) // items_por_pagina + 1 if total_items > 0 else 1
        if st.session_state.pagina_actual_historico > total_paginas:
            st.session_state.pagina_actual_historico = total_paginas
        if st.session_state.pagina_actual_historico < 1:
            st.session_state.pagina_actual_historico = 1
        inicio = (st.session_state.pagina_actual_historico - 1) * items_por_pagina
        fin = inicio + items_por_pagina
        resultados_pagina = resultados_full[inicio:fin]
    else:
        # Paginación en SQLite
        conn = sqlite3.connect(Path(__file__).parent.parent.parent / "analisis.db")
        c = conn.cursor()
        base_where = "WHERE estado = '✅ Completed'"
        params = []
        if filtro_nombre:
            base_where += " AND filename LIKE ?"
            params.append(f"%{filtro_nombre}%")
        if 'fecha_desde' in locals() and fecha_desde:
            base_where += " AND date(created_at) >= ?"
            params.append(fecha_desde.strftime("%Y-%m-%d"))
        if 'fecha_hasta' in locals() and fecha_hasta:
            base_where += " AND date(created_at) <= ?"
            params.append(fecha_hasta.strftime("%Y-%m-%d"))

        c.execute(f"SELECT COUNT(*) FROM analisis {base_where}", params)
        total_items = c.fetchone()[0]
        total_paginas = (total_items - 1) // items_por_pagina + 1 if total_items > 0 else 1
        if st.session_state.pagina_actual_historico > total_paginas:
            st.session_state.pagina_actual_historico = total_paginas
        if st.session_state.pagina_actual_historico < 1:
            st.session_state.pagina_actual_historico = 1

        offset = (st.session_state.pagina_actual_historico - 1) * items_por_pagina
        query = f"""
            SELECT id, filename, created_at
            FROM analisis
            {base_where}
            ORDER BY created_at DESC
            LIMIT ? OFFSET ?
        """
        c.execute(query, params + [items_por_pagina, offset])
        resultados_pagina = c.fetchall()
        conn.close()

    if total_items > 0:
        st.markdown(f"""
            <h3 style='
                color: #374151;
                margin-bottom: 1rem;
                font-size: 1.1rem;
            '>📋 Resultados ({total_items} análisis encontrados)</h3>
        """, unsafe_allow_html=True)
        
        # Paginación superior discreta
        if total_paginas > 1:
            st.markdown(f"""
                <div style='
                    background: #f8fafc;
                    border: 1px solid #e2e8f0;
                    border-radius: 6px;
                    padding: 0.4rem 0.6rem;
                    margin: 0.5rem 0;
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    font-size: 0.75rem;
                    color: #64748b;
                '>
                    <span>📄 Página {st.session_state.pagina_actual_historico} de {total_paginas}</span>
                    <span>📊 {total_items} análisis total</span>
                </div>
            """, unsafe_allow_html=True)
            
            # Navegación con botones más compactos - ocupando todo el ancho
            col_nav1, col_nav2, col_nav3, col_nav4, col_nav5 = st.columns([1, 1, 3, 1, 1])
            
            with col_nav1:
                if st.button("⏪", disabled=st.session_state.pagina_actual_historico == 1, help="Primera página", key="first_top", use_container_width=True):
                    st.session_state.pagina_actual_historico = 1
                    st.rerun()
            
            with col_nav2:
                if st.button("◀", disabled=st.session_state.pagina_actual_historico == 1, help="Página anterior", key="prev_top", use_container_width=True):
                    st.session_state.pagina_actual_historico -= 1
                    st.rerun()
            
            with col_nav3:
                # Mostrar página actual de forma simple
                st.markdown(f"""
                    <div style='
                        text-align: center;
                        color: #64748b;
                        font-size: 0.8rem;
                        font-weight: 500;
                        padding: 0.3rem;
                    '>
                        Página {st.session_state.pagina_actual_historico} de {total_paginas}
                    </div>
                """, unsafe_allow_html=True)
            
            with col_nav4:
                if st.button("▶", disabled=st.session_state.pagina_actual_historico == total_paginas, help="Página siguiente", key="next_top", use_container_width=True):
                    st.session_state.pagina_actual_historico += 1
                    st.rerun()
            
            with col_nav5:
                if st.button("⏩", disabled=st.session_state.pagina_actual_historico == total_paginas, help="Última página", key="last_top", use_container_width=True):
                    st.session_state.pagina_actual_historico = total_paginas
                    st.rerun()
        
        # Mostrar elementos de la página actual
        items_pagina = resultados_pagina
        
        # Lista de análisis
        for analisis_id, filename, created_at in items_pagina:
            # Formatear fecha
            try:
                fecha_obj = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                fecha_formateada = fecha_obj.strftime("%d/%m/%Y %H:%M")
            except:
                fecha_formateada = created_at
            
            filename_corto = filename[:50] + ('...' if len(filename) > 50 else '')
            
            # Card del análisis muy compacto y pulido
            with st.container():
                st.markdown(f"""
                    <div style='
                        background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
                        border: 1px solid #e2e8f0;
                        border-left: 3px solid #16a34a;
                        border-radius: 6px;
                        padding: 0.6rem 0.8rem;
                        margin-bottom: 0.5rem;
                        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
                        transition: all 0.2s ease;
                    '>
                        <div style='
                            display: flex;
                            align-items: center;
                            gap: 0.5rem;
                            margin-bottom: 0.3rem;
                        '>
                            <div style='
                                background: #16a34a;
                                color: white;
                                padding: 0.2rem;
                                border-radius: 4px;
                                font-size: 0.65rem;
                                font-weight: 600;
                            '>✅</div>
                            <h4 style='
                                margin: 0;
                                color: #0f172a;
                                font-size: 0.85rem;
                                font-weight: 600;
                            '>{filename_corto}</h4>
                            <div style='
                                background: #dcfce7;
                                color: #16a34a;
                                padding: 0.1rem 0.4rem;
                                border-radius: 12px;
                                font-size: 0.6rem;
                                font-weight: 500;
                                margin-left: auto;
                            '>COMPLETADO</div>
                        </div>
                        <div style='
                            display: flex;
                            align-items: center;
                            gap: 0.8rem;
                            color: #64748b;
                            font-size: 0.7rem;
                            margin-left: 1.3rem;
                        '>
                            <span>🆔 {analisis_id[:8]}...</span>
                            <span>📅 {fecha_formateada}</span>
                        </div>
                    </div>
                """, unsafe_allow_html=True)
                
                # Botones de acción más compactos - todos del mismo tamaño
                col_action1, col_action2, col_action3 = st.columns([1.5, 1.5, 1.5], vertical_alignment="center")
                
                with col_action1:
                    if st.button("👁️ Ver Detalle", key=f"detalle_hist_{analisis_id}", use_container_width=True, type="primary"):
                        pass
                        mostrar_detalle_dialog(analisis_id, filename)
                
                with col_action2:
                    # Verificar si existe el archivo de progreso y usar caché para generar Word
                    progreso_path = _progreso_path(analisis_id)
                    if progreso_path.exists():
                        try:
                            mtime = os.path.getmtime(progreso_path)
                            word_data = _generar_word_cached(analisis_id, filename, mtime)
                            if word_data:
                                st.download_button(
                                    label="📄 Word MAXAM",
                                    data=word_data,
                                    file_name=f"Analisis_Riesgos_MAXAM_{analisis_id[:8]}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"download_word_hist_{analisis_id}",
                                    use_container_width=True,
                                    help="Descargar informe corporativo MAXAM",
                                    type="primary"
                                )
                            else:
                                st.button("📄 No disponible", disabled=True, 
                                          key=f"no_data_hist_{analisis_id}",
                                          use_container_width=True)
                        except Exception as e:
                            st.button("📄 Error", disabled=True, 
                                      key=f"error_hist_{analisis_id}",
                                      use_container_width=True,
                                      help=f"Error: {str(e)}")
                    else:
                        st.button("📄 No disponible", disabled=True, 
                                  key=f"no_download_hist_{analisis_id}",
                                  use_container_width=True)
                
                with col_action3:
                    if st.button("🔄 Re-análisis", key=f"reanalize_hist_{analisis_id}", 
                               help="Re-analizar documento",
                               use_container_width=True, type="primary"):
                        st.info("🔄 Función de re-análisis en desarrollo")
                
                # Pequeño espaciado entre cards
                st.markdown("<div style='margin-bottom: 0.3rem;'></div>", unsafe_allow_html=True)
        
        # Paginación inferior (repetir si hay muchos elementos)
        if total_paginas > 1:
            st.markdown(f"""
                <div style='
                    background: #f8fafc;
                    border: 1px solid #e2e8f0;
                    border-radius: 6px;
                    padding: 0.4rem 0.6rem;
                    margin: 1rem 0 0.5rem 0;
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    font-size: 0.75rem;
                    color: #64748b;
                '>
                    <span>📄 Página {st.session_state.pagina_actual_historico} de {total_paginas}</span>
                    <span>📊 Mostrando {len(items_pagina)} de {total_items} análisis</span>
                </div>
            """, unsafe_allow_html=True)
            
            # Navegación con botones más compactos - ocupando todo el ancho
            col_nav1, col_nav2, col_nav3, col_nav4, col_nav5 = st.columns([1, 1, 3, 1, 1])
            
            with col_nav1:
                if st.button("⏪", disabled=st.session_state.pagina_actual_historico == 1, help="Primera página", key="first_bottom", use_container_width=True):
                    st.session_state.pagina_actual_historico = 1
                    st.rerun()
            
            with col_nav2:
                if st.button("◀", disabled=st.session_state.pagina_actual_historico == 1, help="Página anterior", key="prev_bottom", use_container_width=True):
                    st.session_state.pagina_actual_historico -= 1
                    st.rerun()
            
            with col_nav3:
                # Mostrar página actual de forma simple
                st.markdown(f"""
                    <div style='
                        text-align: center;
                        color: #64748b;
                        font-size: 0.8rem;
                        font-weight: 500;
                        padding: 0.3rem;
                    '>
                        Página {st.session_state.pagina_actual_historico} de {total_paginas}
                    </div>
                """, unsafe_allow_html=True)
            
            with col_nav4:
                if st.button("▶", disabled=st.session_state.pagina_actual_historico == total_paginas, help="Página siguiente", key="next_bottom", use_container_width=True):
                    st.session_state.pagina_actual_historico += 1
                    st.rerun()
            
            with col_nav5:
                if st.button("⏩", disabled=st.session_state.pagina_actual_historico == total_paginas, help="Última página", key="last_bottom", use_container_width=True):
                    st.session_state.pagina_actual_historico = total_paginas
                    st.rerun()
    
    else:
        # No hay resultados
        filtros_activos = []
        if filtro_nombre:
            filtros_activos.append(f"nombre contiene '{filtro_nombre}'")
        if 'fecha_desde' in locals() and fecha_desde:
            filtros_activos.append(f"desde {fecha_desde}")
        if 'fecha_hasta' in locals() and fecha_hasta:
            filtros_activos.append(f"hasta {fecha_hasta}")
        
        filtros_texto = " • ".join(filtros_activos) if filtros_activos else "Sin filtros aplicados"
        
        st.markdown(f"""
            <div style="
                background: #fef3c7;
                border: 1px dashed #f59e0b;
                border-radius: 8px;
                padding: 1.5rem 1rem;
                text-align: center;
                margin: 1rem 0;
            ">
                <div style="font-size: 2rem; margin-bottom: 0.5rem;">📭</div>
                <h3 style="
                    margin: 0 0 0.25rem 0;
                    color: #92400e;
                    font-size: 1rem;
                    font-weight: 600;
                ">No se encontraron análisis</h3>
                <p style="
                    margin: 0;
                    color: #a16207;
                    font-size: 0.8rem;
                ">No hay análisis que coincidan con los criterios: {filtros_texto}</p>
            </div>
        """, unsafe_allow_html=True)

def generar_documento_word_corporativo(progreso_data, preguntas, resultados, filename):
    """
    Genera un documento Word corporativo de MAXAM con el formato estándar de la empresa
    """
    try:
        # Crear un nuevo documento
        doc = Document()
        
        # ===============================
        # ENCABEZADO CORPORATIVO
        # ===============================
        
        # Agregar logo (si existe)
        logo_path = Path("src/images/maxam_logo.png")
        if logo_path.exists():
            try:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.add_picture(str(logo_path), width=Inches(2.5))
            except:
                # Si falla el logo, agregar texto del encabezado
                header = doc.add_heading('MAXAM', level=1)
                header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Título principal
        title = doc.add_heading('ANÁLISIS DE RIESGOS CONTRACTUALES', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del documento
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        info_table.autofit = True
        
        # Configurar encabezados de la tabla de información
        info_cells = [
            ('Documento Analizado:', filename),
            ('Fecha de Análisis:', datetime.now().strftime("%d/%m/%Y %H:%M")),
            ('ID de Análisis:', progreso_data.get('id', 'N/A')[:8]),
            ('Total de Preguntas:', str(len(resultados))),
            ('Sistema:', 'AI Risk Analyzer - MAXAM'),
            ('Versión:', '1.0')
        ]
        
        for i, (label, value) in enumerate(info_cells):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
            # Hacer la primera columna en negrita
            info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        
        # Espacio
        doc.add_paragraph()
        
        # ===============================
        # RESUMEN EJECUTIVO
        # ===============================
        
        doc.add_heading('RESUMEN EJECUTIVO', level=2)
        
        # Contar riesgos por nivel
        riesgos = {'Alto': 0, 'Medio': 0, 'Bajo': 0, 'Sin evaluar': 0}
        for resultado in resultados:
            riesgo = resultado.get('Riesgo', 'Sin evaluar')
            if riesgo in riesgos:
                riesgos[riesgo] += 1
            else:
                riesgos['Sin evaluar'] += 1
        
        total_preguntas = len(resultados)
        
        # Tabla de resumen de riesgos
        resumen_table = doc.add_table(rows=5, cols=3)
        resumen_table.style = 'Table Grid'
        resumen_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados
        headers = resumen_table.rows[0].cells
        headers[0].text = 'Nivel de Riesgo'
        headers[1].text = 'Cantidad'
        headers[2].text = 'Porcentaje'
        
        # Hacer encabezados en negrita
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Datos de riesgo
        risk_rows = [
            ('🔴 Alto', riesgos['Alto'], f"{riesgos['Alto']/total_preguntas*100:.1f}%"),
            ('🟡 Medio', riesgos['Medio'], f"{riesgos['Medio']/total_preguntas*100:.1f}%"),
            ('🟢 Bajo', riesgos['Bajo'], f"{riesgos['Bajo']/total_preguntas*100:.1f}%"),
            ('⚪ Sin evaluar', riesgos['Sin evaluar'], f"{riesgos['Sin evaluar']/total_preguntas*100:.1f}%")
        ]
        
        for i, (nivel, cantidad, porcentaje) in enumerate(risk_rows, 1):
            row_cells = resumen_table.rows[i].cells
            row_cells[0].text = nivel
            row_cells[1].text = str(cantidad)
            row_cells[2].text = porcentaje
        
        # Párrafo de conclusión
        doc.add_paragraph()
        conclusion_text = f"""
        El análisis del documento "{filename}" ha identificado un total de {total_preguntas} aspectos evaluados. 
        Se han detectado {riesgos['Alto']} riesgos de nivel alto, {riesgos['Medio']} de nivel medio y {riesgos['Bajo']} de nivel bajo.
        """
        conclusion = doc.add_paragraph(conclusion_text.strip())
        
        # ===============================
        # ANÁLISIS DETALLADO
        # ===============================
        
        doc.add_page_break()
        doc.add_heading('ANÁLISIS DETALLADO POR SECCIONES', level=2)
        
        # Agrupar preguntas por sección
        secciones = {}
        for i, (pregunta_original, resultado) in enumerate(zip(preguntas, resultados)):
            # Usar la pregunta y sección más reciente (reanalizada si existe)
            if resultado.get('reanalizado_en') and resultado.get('Pregunta'):
                pregunta_texto = resultado.get('Pregunta')
                seccion = resultado.get('Sección', pregunta_original.get('Sección', 'Sin sección'))
            else:
                pregunta_texto = pregunta_original.get('Pregunta', 'N/A')
                seccion = pregunta_original.get('Sección', 'Sin sección')
            
            if seccion not in secciones:
                secciones[seccion] = []
            
            secciones[seccion].append({
                'numero': i + 1,
                'pregunta': pregunta_texto,
                'respuesta': resultado.get('Respuesta', ''),
                'riesgo': resultado.get('Riesgo', 'Sin evaluar'),
                'reanalizada': bool(resultado.get('reanalizado_en')),
                'fecha_reanalisis': resultado.get('reanalizado_en', '')
            })
        
        # Procesar cada sección
        for seccion, items in secciones.items():
            # Título de sección
            doc.add_heading(f'SECCIÓN: {seccion.upper()}', level=3)
            
            # Procesar cada pregunta de la sección
            for item in items:
                # Número y pregunta
                pregunta_p = doc.add_paragraph()
                pregunta_run = pregunta_p.add_run(f"P{item['numero']}: ")
                pregunta_run.font.bold = True
                pregunta_run.font.size = Pt(11)
                
                pregunta_texto = item['pregunta']
                
                # Si la pregunta fue reanalizada, agregar indicador
                if item['reanalizada']:
                    pregunta_texto += " 🔄 (Reanalizada)"
                
                pregunta_text_run = pregunta_p.add_run(pregunta_texto)
                pregunta_text_run.font.size = Pt(10)
                
                # Indicador de reanalizada si aplica
                if item['reanalizada']:
                    reanalisis_p = doc.add_paragraph()
                    reanalisis_run = reanalisis_p.add_run("🔄 PREGUNTA REANALIZADA")
                    reanalisis_run.font.bold = True
                    reanalisis_run.font.color.rgb = RGBColor(133, 77, 14)  # Color amarillo/dorado
                    reanalisis_run.font.size = Pt(9)
                    
                    if item['fecha_reanalisis']:
                        fecha_run = reanalisis_p.add_run(f" - {item['fecha_reanalisis']}")
                        fecha_run.font.size = Pt(8)
                        fecha_run.font.color.rgb = RGBColor(107, 114, 128)  # Gris
                
                # Nivel de riesgo
                riesgo_p = doc.add_paragraph()
                riesgo_label_run = riesgo_p.add_run("NIVEL DE RIESGO: ")
                riesgo_label_run.font.bold = True
                riesgo_label_run.font.size = Pt(10)
                
                riesgo_value_run = riesgo_p.add_run(item['riesgo'])
                riesgo_value_run.font.bold = True
                riesgo_value_run.font.size = Pt(10)
                
                # Colorear según el riesgo
                if item['riesgo'] == 'Alto':
                    riesgo_value_run.font.color.rgb = RGBColor(220, 38, 38)  # Rojo
                elif item['riesgo'] == 'Medio':
                    riesgo_value_run.font.color.rgb = RGBColor(245, 158, 11)  # Amarillo
                elif item['riesgo'] == 'Bajo':
                    riesgo_value_run.font.color.rgb = RGBColor(22, 163, 74)  # Verde
                else:
                    riesgo_value_run.font.color.rgb = RGBColor(107, 114, 128)  # Gris
                
                # Respuesta con formato markdown procesado
                respuesta_label_p = doc.add_paragraph()
                respuesta_label_run = respuesta_label_p.add_run("ANÁLISIS:")
                respuesta_label_run.font.bold = True
                respuesta_label_run.font.size = Pt(10)
                
                # Procesar markdown y agregar al documento
                procesar_markdown_a_word(doc, item['respuesta'])
                
                # Línea separadora
                doc.add_paragraph("_" * 80)
        
        # ===============================
        # PIE DE PÁGINA
        # ===============================
        
        doc.add_page_break()
        
        # Información adicional
        doc.add_heading('INFORMACIÓN ADICIONAL', level=2)
        
        info_adicional = f"""
        Este documento ha sido generado automáticamente por el sistema AI Risk Analyzer de MAXAM.
        
        Fecha de generación: {datetime.now().strftime("%d de %B de %Y a las %H:%M")}
        
        El análisis se basa en inteligencia artificial y debe ser revisado por expertos legales y técnicos 
        antes de tomar decisiones contractuales definitivas.
        
        Para consultas sobre este análisis, contacte al equipo de gestión de riesgos de MAXAM.
        """
        
        doc.add_paragraph(info_adicional.strip())
        
        # Firma corporativa
        doc.add_paragraph()
        firma_p = doc.add_paragraph()
        firma_run = firma_p.add_run("MAXAM - Análisis de Riesgos Contractuales")
        firma_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        firma_run.font.bold = True
        firma_run.font.size = Pt(12)
        
        # Guardar en memoria
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
        
    except Exception as e:
        # Usar print si streamlit no está disponible
        try:
            import streamlit as st
            st.error(f"Error al generar documento Word: {str(e)}")
        except:
            print(f"Error al generar documento Word: {str(e)}")
        return None

def exportar_historico_word():
    """Exporta el histórico completo a Word corporativo MAXAM"""
    try:
        import sqlite3
        
        conn = sqlite3.connect(Path(__file__).parent.parent.parent / "analisis.db")
        
        # Obtener datos del histórico
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id, filename, estado, created_at 
            FROM analisis 
            ORDER BY created_at DESC
        """)
        
        analisis_list = cursor.fetchall()
        conn.close()
        
        if not analisis_list:
            try:
                import streamlit as st
                st.warning("No hay análisis en el histórico para exportar")
            except:
                print("Warning: No hay análisis en el histórico para exportar")
            return None
        
        # Crear documento Word
        doc = Document()
        
        # Encabezado corporativo
        title = doc.add_heading('HISTÓRICO DE ANÁLISIS DE RIESGOS - MAXAM', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del reporte
        info_p = doc.add_paragraph()
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_p.add_run(f"Reporte generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}")
        info_run.font.size = Pt(10)
        info_run.font.color.rgb = RGBColor(107, 114, 128)
        
        doc.add_paragraph()
        
        # Resumen
        doc.add_heading('RESUMEN', level=2)
        
        total_analisis = len(analisis_list)
        completados = sum(1 for a in analisis_list if a[2] == "✅ Completed")
        en_proceso = sum(1 for a in analisis_list if "🔄" in a[2])
        errores = sum(1 for a in analisis_list if "❌" in a[2])
        
        resumen_text = f"""
        Total de análisis registrados: {total_analisis}
        • Análisis completados: {completados}
        • Análisis en proceso: {en_proceso} 
        • Análisis con errores: {errores}
        """
        
        doc.add_paragraph(resumen_text.strip())
        
        # Tabla detallada
        doc.add_heading('DETALLE DE ANÁLISIS', level=2)
        
        # Crear tabla
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = True
        
        # Encabezados
        headers = table.rows[0].cells
        headers[0].text = 'ID'
        headers[1].text = 'Documento'
        headers[2].text = 'Estado'
        headers[3].text = 'Fecha'
        
        # Hacer encabezados en negrita
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Agregar datos
        for analisis in analisis_list:
            row_cells = table.add_row().cells
            row_cells[0].text = analisis[0][:8] + "..."  # ID truncado
            row_cells[1].text = analisis[1][:50] + "..." if len(analisis[1]) > 50 else analisis[1]  # Filename truncado
            row_cells[2].text = analisis[2]  # Estado
            
            # Formatear fecha
            try:
                fecha_dt = datetime.fromisoformat(analisis[3].replace('Z', '+00:00'))
                fecha_formateada = fecha_dt.strftime('%d/%m/%Y %H:%M')
            except:
                fecha_formateada = analisis[3]
            
            row_cells[3].text = fecha_formateada
        
        # Pie de página
        doc.add_page_break()
        
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run("MAXAM - Sistema de Análisis de Riesgos Contractuales")
        footer_run.font.bold = True
        footer_run.font.size = Pt(12)
        
        # Guardar en memoria
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
        
    except Exception as e:
        try:
            import streamlit as st
            st.error(f"Error al exportar histórico a Word: {str(e)}")
        except:
            print(f"Error al exportar histórico a Word: {str(e)}")
        return None

def procesar_markdown_a_word(doc, texto_markdown):
    """
    Convierte texto markdown a formato Word nativo con bullets, negrita, etc.
    """
    if not texto_markdown:
        return
    
    # Dividir el texto en líneas para procesarlo
    lineas = texto_markdown.split('\n')
    
    i = 0
    while i < len(lineas):
        linea = lineas[i].strip()
        
        if not linea:
            # Línea vacía - agregar espacio
            doc.add_paragraph()
            i += 1
            continue
        
        # Detectar listas con bullets (-, *, +)
        if re.match(r'^[-*+]\s+', linea):
            # Procesar lista de bullets
            lista_items = []
            while i < len(lineas) and re.match(r'^[-*+]\s+', lineas[i].strip()):
                item_text = re.sub(r'^[-*+]\s+', '', lineas[i].strip())
                lista_items.append(item_text)
                i += 1
            
            # Agregar cada item como bullet point
            for item in lista_items:
                p = doc.add_paragraph(style='List Bullet')
                procesar_texto_con_formato(p, item)
            
            continue
        
        # Detectar listas numeradas (1., 2., etc.)
        elif re.match(r'^\d+\.\s+', linea):
            # Procesar lista numerada
            lista_items = []
            while i < len(lineas) and re.match(r'^\d+\.\s+', lineas[i].strip()):
                item_text = re.sub(r'^\d+\.\s+', '', lineas[i].strip())
                lista_items.append(item_text)
                i += 1
            
            # Agregar cada item como numbered list
            for item in lista_items:
                p = doc.add_paragraph(style='List Number')
                procesar_texto_con_formato(p, item)
            
            continue
        
        # Detectar encabezados (####, ###, ##, #)
        elif re.match(r'^#+\s+', linea):
            nivel = len(re.match(r'^#+', linea).group())
            texto_header = re.sub(r'^#+\s+', '', linea)
            
            # Agregar encabezado según el nivel
            if nivel == 1:
                h = doc.add_heading(texto_header, level=3)
            elif nivel == 2:
                h = doc.add_heading(texto_header, level=4)
            elif nivel == 3:
                h = doc.add_heading(texto_header, level=5)
            else:
                # Para niveles 4+ usar un párrafo con formato destacado
                p = doc.add_paragraph()
                run = p.add_run(texto_header)
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(44, 62, 80)  # Color azul oscuro
            
            i += 1
            continue
        
        # Detectar bloques de código (```texto```)
        elif linea.startswith('```'):
            # Procesar bloque de código
            codigo_lineas = []
            i += 1  # Saltar la línea de apertura
            
            while i < len(lineas) and not lineas[i].strip().startswith('```'):
                codigo_lineas.append(lineas[i])
                i += 1
            
            if i < len(lineas) and lineas[i].strip().startswith('```'):
                i += 1  # Saltar la línea de cierre
            
            # Agregar código como párrafo con formato especial
            if codigo_lineas:
                codigo_texto = '\n'.join(codigo_lineas)
                p = doc.add_paragraph()
                run = p.add_run(codigo_texto)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            
            continue
        
        # Línea normal - procesar con formato inline
        else:
            p = doc.add_paragraph()
            procesar_texto_con_formato(p, linea)
            i += 1

def procesar_texto_con_formato(paragraph, texto):
    """
    Procesa texto con formato markdown inline (negrita, cursiva, código)
    y lo agrega al párrafo de Word
    """
    if not texto:
        return
    
    # Patrones para formato inline
    # **texto** o __texto__ para negrita
    # *texto* o _texto_ para cursiva
    # `texto` para código inline
    
    # Crear una lista de segmentos con su formato
    segmentos = []
    texto_restante = texto
    
    # Procesar negrita (**texto**)
    while '**' in texto_restante:
        antes, resto = texto_restante.split('**', 1)
        if '**' in resto:
            negrita, despues = resto.split('**', 1)
            
            if antes:
                segmentos.append(('normal', antes))
            segmentos.append(('bold', negrita))
            texto_restante = despues
        else:
            # No hay cierre, tratar como texto normal
            segmentos.append(('normal', texto_restante))
            break
    
    if texto_restante and not segmentos:
        segmentos.append(('normal', texto_restante))
    elif texto_restante:
        segmentos.append(('normal', texto_restante))
    
    # Segunda pasada para cursiva en cada segmento normal
    segmentos_finales = []
    for tipo, texto_seg in segmentos:
        if tipo == 'bold':
            segmentos_finales.append((tipo, texto_seg))
        else:
            # Procesar cursiva en texto normal
            texto_temp = texto_seg
            while '*' in texto_temp and texto_temp.count('*') >= 2:
                antes, resto = texto_temp.split('*', 1)
                if '*' in resto:
                    cursiva, despues = resto.split('*', 1)
                    
                    if antes:
                        segmentos_finales.append(('normal', antes))
                    segmentos_finales.append(('italic', cursiva))
                    texto_temp = despues
                else:
                    segmentos_finales.append(('normal', texto_temp))
                    break
            
            if texto_temp:
                segmentos_finales.append(('normal', texto_temp))
    
    # Tercera pasada para código inline en cada segmento normal
    segmentos_codigo = []
    for tipo, texto_seg in segmentos_finales:
        if tipo in ['bold', 'italic']:
            segmentos_codigo.append((tipo, texto_seg))
        else:
            # Procesar código inline en texto normal
            texto_temp = texto_seg
            while '`' in texto_temp:
                antes, resto = texto_temp.split('`', 1)
                if '`' in resto:
                    codigo, despues = resto.split('`', 1)
                    
                    if antes:
                        segmentos_codigo.append(('normal', antes))
                    segmentos_codigo.append(('code', codigo))
                    texto_temp = despues
                else:
                    segmentos_codigo.append(('normal', texto_temp))
                    break
            
            if texto_temp:
                segmentos_codigo.append(('normal', texto_temp))
    
    # Agregar los segmentos al párrafo con su formato correspondiente
    for tipo, texto_seg in segmentos_codigo:
        if not texto_seg:
            continue
            
        run = paragraph.add_run(texto_seg)
        run.font.size = Pt(10)
        
        if tipo == 'bold':
            run.font.bold = True
        elif tipo == 'italic':
            run.font.italic = True
        elif tipo == 'code':
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            # No podemos agregar fondo fácilmente al run, pero sí cambiar la fuente
